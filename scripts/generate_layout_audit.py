"""Generate the table-by-table clinical/statistical layout audit deliverables."""

from __future__ import annotations

import argparse
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

from tflshell.data.definitions import build_catalog

RETIRED_TABLES = {
    "T14.3.1.8": (
        "TEAEs Occurring in >=5% of Subjects by Preferred Term",
        "Duplicates the thresholded SOC/PT summary in T14.3.1.2.",
    ),
    "T14.3.1.9": (
        "TEAEs by Cycle",
        "AEs are recorded by event date; a generic cycle summary is not clinically standard.",
    ),
    "T14.3.1.16": (
        "TEAEs by SOC and Maximum CTCAE Grade",
        "Duplicates the consolidated SOC/PT/maximum-grade table T14.3.1.3.",
    ),
    "T14.3.1.17": (
        "TEAEs by Preferred Term — Full Frequency Listing (All PTs)",
        "This is a SOC/PT summary rather than a patient listing and duplicates T14.3.1.2.",
    ),
    "T14.3.1.18": (
        "TEAEs — Maximum Severity by Relationship to Study Drug",
        "The relationship-by-grade cross-layout is nonstandard; relatedness and grade are reviewed separately.",
    ),
    "T14.3.1.21": (
        "TEAEs by Preferred Term and Maximum CTCAE Grade — Group 2",
        "Group-specific duplication was replaced by one all-treatment-group grade table.",
    ),
    "T14.3.1.22": (
        "TEAEs by Cycle — Detailed On-Treatment Period Analysis",
        "A generic AE-by-cycle layout is not supported without an explicit protocol/SAP rationale.",
    ),
    "T14.3.1.29": (
        "Treatment-Emergent SAEs by PT and ICH E2A Criterion — Group 1",
        "Group-specific multi-response criterion cross-tab is not a standard CSR summary; retain in the SAE listing if required.",
    ),
    "T14.3.1.30": (
        "Recurrent TEAEs — Subjects with Same PT >=2 Occurrences",
        "The original columns mixed subject frequency, event frequency, and group-specific counts in a non-comparable layout.",
    ),
}

RESTRUCTURED_IDS = {
    "T14.3.1.3",
    "T14.3.1.4",
    "T14.3.1.13",
    "T14.3.1.19",
    "T14.3.1.20",
    "T14.3.1.25",
    "T14.3.1.26",
    "T14.3.1.27",
    "T14.3.1.28",
    "T14.3.1.31",
    "T14.4.4",
    "T14.4.5",
    "T14.4.6",
}


def table_category(item_id: str, title: str) -> str:
    lowered = title.lower()
    if item_id.startswith("T14.1"):
        return "14.1 Subject Accounting / Baseline"
    if item_id.startswith("T14.2"):
        if any(token in lowered for token in ("survival", "time to", "duration of response")):
            return "14.2 Time-to-Event Efficacy"
        if any(token in lowered for token in ("response", "responder", "odds ratio")):
            return "14.2 Binary / Response Efficacy"
        return "14.2 Continuous / Other Efficacy"
    if item_id.startswith("T14.3.1"):
        return "14.3 Adverse Events"
    if item_id.startswith("T14.3.2"):
        return "14.3 Deaths and Subject-Level Safety Summary"
    if item_id.startswith("T14.3.3"):
        return "14.3 Laboratory Safety"
    if item_id.startswith("T14.3.4"):
        return "14.3 Vital Signs / ECG"
    if any(token in lowered for token in ("pharmacokinetic", "pk ", "concentration", "auc")):
        return "14.4 Pharmacokinetics"
    return "14.4 Biomarker / PD / PRO / Immunogenicity"


def clinical_judgement(item) -> str:
    title = item.title.lower()
    first_column = item.placeholder_columns[0] if item.placeholder_columns else ""
    if item.id == "T14.3.1.3":
        return "Appropriate after revision: SOC > PT > maximum grade is a row hierarchy; all treatment groups are comparable in columns."
    if item.id.startswith("T14.3.1"):
        if any(token in title for token in ("late-onset", "time window", "follow-up", "infusion-related", "immune-related")):
            return "Retain conditionally: the event concept/window must be protocol- or SAP-defined; no generic cycle assumption is used."
        return "Appropriate: subject/event meaning is stated, SOC/PT or event category remains structural, and treatment groups are comparable."
    if item.layout_profile == "model-comparison":
        return "Appropriate: treatment estimates and between-group comparison occupy independent column groups; model quantities remain SAP-dependent."
    if "by visit" in title:
        if first_column.startswith("Visit /") or first_column == "Visit":
            return "Appropriate: Visit is the highest row hierarchy, followed by endpoint/parameter and statistic."
        if "Parameter" in first_column:
            return "Appropriate for a multi-parameter safety review: Parameter > Visit > Statistic is retained to support programming and medical review."
        return "Appropriate: visit/timepoint is explicit and not hidden inside a treatment-group result cell."
    if any(token in title for token in ("shift table", "worst post-baseline")):
        return "Appropriate: baseline and post-baseline category dimensions are structural; counts/percentages remain treatment-group results."
    if any(token in title for token in ("survival", "time to", "duration")):
        return "Appropriate if the SAP supplies time origin, event, censoring, analysis set, and estimator definitions."
    return "Appropriate for the stated reporting purpose; no unsupported model result or artificial grouping was added."


def programming_note(item) -> str:
    first_column = item.placeholder_columns[0] if item.placeholder_columns else ""
    if item.layout_profile == "model-comparison":
        return "Create treatment-estimate rows and comparison rows separately, then join on endpoint/visit/statistic and render grouped headers."
    if item.id == "T14.3.1.3":
        return "Derive each subject's maximum grade within SOC/PT; produce a long SOC/PT/grade shell and transpose treatment groups only."
    if "Visit" in first_column or "Timepoint" in first_column:
        return "Keep visit/parameter/statistic as ordered keys in long form; transpose only treatment-group result values."
    if "System Organ Class" in first_column or "Preferred Term" in first_column:
        return "Use controlled SOC/PT sort keys and indentation flags; treatment columns are produced from the same subject/event denominator rules."
    return "Use the first column as the structural key; create one result variable per displayed treatment/statistic column."


def footnote_status(item) -> str:
    notes = item.footnote_text()
    abbreviation_ok = (
        not any(note.startswith("Abbreviations:") for note in item.footnotes)
        or any(" = " in note for note in item.footnotes if note.startswith("Abbreviations:"))
    )
    version_ok = all(
        sum(name in note and version in note for note in notes) == 1
        for name, version in item.dictionary_versions.items()
    )
    return (
        "Pass — abbreviations/definitions are controlled; coding/grading versions occur once."
        if abbreviation_ok and version_ok
        else "Review required — footnote contract check did not pass."
    )


def audit_rows() -> list[dict[str, str]]:
    catalog = build_catalog()
    rows: list[dict[str, str]] = []
    for item in catalog.tables():
        category = table_category(item.id, item.title)
        conditional = item.coverage_summary.startswith("Conditional")
        if item.id in RESTRUCTURED_IDS:
            disposition = "Retained — restructured"
        elif conditional:
            disposition = "Retained — conditional"
        else:
            disposition = "Retained — acceptable"
        rows.append(
            {
                "TFL ID": item.id,
                "Title": item.title,
                "Category": category,
                "Population": item.population,
                "Row hierarchy": item.placeholder_columns[0] if item.placeholder_columns else "",
                "Column layout": " | ".join(column.replace("\n", " / ") for column in item.placeholder_columns[1:]),
                "Clinical/statistical judgement": clinical_judgement(item),
                "Programming implementation": programming_note(item),
                "Footnote/statistical definition": footnote_status(item),
                "Disposition": disposition,
                "Condition / rationale": (
                    "Use only when required by protocol/SAP and when the displayed population/denominator is defined."
                    if conditional
                    else "No unresolved layout issue identified in this review."
                ),
            }
        )
    for item_id, (title, reason) in RETIRED_TABLES.items():
        rows.append(
            {
                "TFL ID": item_id,
                "Title": title,
                "Category": "14.3 Adverse Events",
                "Population": "Safety Population",
                "Row hierarchy": "Retired",
                "Column layout": "Retired",
                "Clinical/statistical judgement": "Not retained in the governed catalog.",
                "Programming implementation": "Do not program as a standalone summary; use the referenced retained table/listing when applicable.",
                "Footnote/statistical definition": "Not applicable after retirement.",
                "Disposition": "Retired",
                "Condition / rationale": reason,
            }
        )
    return sorted(rows, key=lambda row: [int(part) for part in row["TFL ID"][1:].split(".")])


def style_sheet(ws, widths: list[int], disposition_column: int | None = None) -> None:
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions
    ws.sheet_view.showGridLines = False
    for cell in ws[1]:
        cell.fill = PatternFill("solid", fgColor="1F4E78")
        cell.font = Font(color="FFFFFF", bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
        disposition = row[disposition_column - 1].value if disposition_column else None
        if disposition == "Retired":
            fill = PatternFill("solid", fgColor="F4CCCC")
        elif disposition == "Retained — restructured":
            fill = PatternFill("solid", fgColor="FFF2CC")
        elif disposition == "Retained — conditional":
            fill = PatternFill("solid", fgColor="D9EAD3")
        else:
            fill = None
        if fill:
            for cell in row:
                cell.fill = fill
    for index, width in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(index)].width = width
    ws.page_setup.orientation = "landscape"
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.sheet_properties.pageSetUpPr.fitToPage = True


def build_workbook(rows: list[dict[str, str]], path: Path) -> None:
    wb = Workbook()
    summary = wb.active
    summary.title = "Review Summary"
    summary.sheet_view.showGridLines = False
    summary["A1"] = "TFL Table Layout Audit — R3"
    summary["A1"].font = Font(size=16, bold=True, color="1F4E78")
    summary.merge_cells("A1:D1")
    summary.append([])
    summary.append(["Scope", "Count", "Interpretation", "Review basis"])
    counts = Counter(row["Disposition"] for row in rows)
    summary.append(["Original tables reviewed", len(rows), "Complete table-level review", "Clinical/statistical practice + programming feasibility + information clarity"])
    summary.append(["Retained tables", len(rows) - counts["Retired"], "Current governed catalog", "Includes acceptable, conditional, and restructured tables"])
    summary.append(["Retired tables", counts["Retired"], "Removed from governed catalog", "Duplicate, generic AE-by-cycle, or nonstandard layout"])
    summary.append(["Restructured retained tables", counts["Retained — restructured"], "Material layout correction", "AE hierarchy/timing/outcome or by-visit/model comparison"])
    summary.append([])
    summary.append(["Key standard", "Decision", "Programming implication", "Clinical/statistical rationale"])
    standards = [
        ("AE grade", "Grade is a row hierarchy under SOC/PT, not a result column.", "Derive maximum grade in long form before transposition.", "Keeps treatment groups directly comparable."),
        ("AE time", "Do not summarize general AEs by cycle.", "Use event dates and protocol-defined windows only.", "AE is event-time data; cycle grouping requires explicit rationale."),
        ("By visit", "Visit is highest for a single endpoint; multi-parameter safety may use Parameter > Visit.", "Keep structural keys in rows; transpose only results.", "Balances medical review and implementation feasibility."),
        ("Model comparison", "Treatment estimates and comparisons use separate column groups.", "Join estimate and contrast outputs by endpoint/visit/statistic.", "Avoids placing a comparison under one treatment group."),
        ("Footnotes", "One abbreviation line, one statistical-definition line when needed, one coding/grading version line.", "Generate controlled notes after table content is finalized.", "Prevents duplicated MedDRA/CTCAE notes and undefined statistics."),
    ]
    for standard in standards:
        summary.append(standard)
    for cell in summary[3]:
        cell.fill = PatternFill("solid", fgColor="1F4E78")
        cell.font = Font(color="FFFFFF", bold=True)
    for cell in summary[9]:
        cell.fill = PatternFill("solid", fgColor="1F4E78")
        cell.font = Font(color="FFFFFF", bold=True)
    for row in summary.iter_rows():
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    for column, width in zip("ABCD", (28, 54, 52, 58), strict=True):
        summary.column_dimensions[column].width = width
    summary.freeze_panes = "A3"

    detail = wb.create_sheet("Table-by-Table Audit")
    headers = list(rows[0])
    detail.append(headers)
    for row in rows:
        detail.append([row[header] for header in headers])
    style_sheet(detail, [15, 48, 30, 28, 34, 55, 70, 70, 52, 26, 66], disposition_column=10)

    retired = wb.create_sheet("Retired Tables")
    retired_rows = [row for row in rows if row["Disposition"] == "Retired"]
    retired_headers = ["TFL ID", "Title", "Condition / rationale", "Programming implementation"]
    retired.append(retired_headers)
    for row in retired_rows:
        retired.append([row[header] for header in retired_headers])
    style_sheet(retired, [16, 55, 90, 80])

    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(path)


def add_docx_table(document: Document, headers: list[str], data: list[list[str]], widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        cell._tc.get_or_add_tcPr().append(_cell_shading("1F4E78"))
    for record in data:
        cells = table.add_row().cells
        for index, value in enumerate(record):
            cells[index].text = str(value)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)
    return table


def _cell_shading(fill: str):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    return shading


def build_report(rows: list[dict[str, str]], path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TFL Table Layout Clinical/Statistical Review Report — R3")
    run.bold = True
    run.font.size = Pt(16)
    doc.add_paragraph(
        "Scope: all 140 tables present at the start of this review. The review evaluates clinical/statistical reporting practice, programming feasibility, and information clarity; it does not add unsupported estimands, models, or metrics."
    )

    counts = Counter(row["Disposition"] for row in rows)
    doc.add_heading("1. Review outcome", level=1)
    add_docx_table(
        doc,
        ["Reviewed", "Retained", "Retired", "Restructured"],
        [[len(rows), len(rows) - counts["Retired"], counts["Retired"], counts["Retained — restructured"]]],
        [1.3, 1.3, 1.3, 1.5],
    )
    doc.add_paragraph(
        "The governed catalog now contains 131 tables. Nine AE tables were retired because they were duplicate summaries, generic AE-by-cycle displays, group-specific duplicates, or nonstandard cross-layouts."
    )

    doc.add_heading("2. Review points by table category", level=1)
    category_points = [
        ("14.1 Subject accounting, disposition, exposure, medication, and baseline", "Keep population/denominator and category definitions explicit. Structural classifications such as ATC level, country/site, and reason remain left-aligned row dimensions; treatment results remain comparable columns."),
        ("14.2 Continuous and by-visit efficacy", "For a single endpoint, Visit is the highest row hierarchy. Observed values, model-based treatment estimates, and between-group comparisons are visually separated. Model quantities appear only where the SAP supports them."),
        ("14.2 Binary/response and time-to-event efficacy", "Response definitions and denominators must be explicit. Time origin, event, censoring, estimand, and estimator definitions remain SAP-controlled; the shell does not manufacture them."),
        ("14.3 Adverse events", "SOC/PT is the principal medical hierarchy; maximum CTCAE grade is a row group. General AE is not summarized by cycle. AESI, IRR, late-onset, onset-window, and follow-up tables are conditional on protocol/SAP definitions."),
        ("14.3 Laboratory, vital signs, and ECG", "Multi-parameter medical review may use Parameter > Visit > Statistic. Single-parameter by-visit endpoints use Visit first. Shift and abnormality tables keep category dimensions structural and treatment results in columns."),
        ("14.4 PK, PD, biomarkers, PRO, and immunogenicity", "Actual scheduled timepoint/cycle may be retained where intrinsic to sampling. By-visit endpoints use Visit first; model comparisons are independent from treatment estimates; assay- or design-specific tables remain conditional."),
    ]
    for heading, text in category_points:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(f"{heading}: ").bold = True
        paragraph.add_run(text)

    doc.add_heading("3. Material corrections", level=1)
    changes = [
        ["Footnotes", "Removed repeated inline MedDRA/CTCAE versions; each table now produces one controlled coding/grading version note."],
        ["AE maximum grade", "T14.3.1.3 now compares all treatment groups and displays grade beneath SOC/PT as rows."],
        ["AE timing/outcome", "T14.3.1.19/.20/.26 place statistic/outcome in rows and treatment groups in columns; protocol-defined windows replace arbitrary cutoffs."],
        ["AE scope", "T14.3.1.25 removes the assumption that first infusion equals Cycle 1; .27/.28/.31 use protocol-defined late/follow-up windows."],
        ["By visit", "T14.4.4/.5/.6 place Visit at the highest level; T14.4.5 uses separate treatment-estimate and comparison column groups."],
        ["Semantic columns", "Legacy ‘Visit — Statistic’ row labels are expanded into real Visit/Timepoint/Statistic fields before rendering, preventing placeholder shifts."],
    ]
    add_docx_table(doc, ["Area", "Correction"], changes, [1.7, 7.2])

    doc.add_heading("4. Retired tables", level=1)
    retired = [
        [row["TFL ID"], row["Title"], row["Condition / rationale"]]
        for row in rows
        if row["Disposition"] == "Retired"
    ]
    add_docx_table(doc, ["TFL ID", "Title", "Reason"], retired, [1.25, 3.1, 4.6])

    doc.add_heading("5. Residual conditions and risks", level=1)
    for text in (
        "Protocol/SAP-specific thresholds, visits, follow-up windows, model covariates, estimands, missing-data handling, and multiplicity rules remain placeholders until study documents are available.",
        "Cycle/timepoint remains valid for DLT, infusion scheduling, PK/PD sampling, or another explicitly scheduled assessment; the AE rule is not applied mechanically to those tables.",
        "Wide tables may still require landscape orientation, continued headers, or controlled column pruning after the actual number of treatment groups and visits is known.",
        "The detailed workbook is the table-level review record; this report summarizes the decisions and does not replace study-specific SAP review.",
    ):
        doc.add_paragraph(text, style="List Bullet")

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def verify_workbook(path: Path, expected_rows: int) -> None:
    wb = load_workbook(path, read_only=False, data_only=False)
    assert wb.sheetnames == ["Review Summary", "Table-by-Table Audit", "Retired Tables"]
    detail = wb["Table-by-Table Audit"]
    assert detail.max_row - 1 == expected_rows
    assert detail.auto_filter.ref == detail.dimensions
    assert detail.freeze_panes == "A2"
    assert wb["Retired Tables"].max_row - 1 == len(RETIRED_TABLES)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output-dir", default="output")
    parser.add_argument("--suffix", default="R3")
    args = parser.parse_args()
    output_dir = Path(args.output_dir)
    rows = audit_rows()
    assert len(rows) == 140, f"Expected 140 reviewed tables, found {len(rows)}"
    workbook_path = output_dir / f"TFL_Table_Layout_Audit_{args.suffix}.xlsx"
    report_path = output_dir / f"TFL_Table_Layout_Audit_Report_{args.suffix}.docx"
    build_workbook(rows, workbook_path)
    build_report(rows, report_path)
    verify_workbook(workbook_path, len(rows))
    print(workbook_path)
    print(report_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
