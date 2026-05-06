"""Word Table of Contents field and bookmark utilities."""

from docx.shared import Pt

from tflshell import config
from tflshell.docx_utils.xml_helpers import insert_field


def insert_toc_field(doc, heading_depth=3):
    """Insert a Word-native Table of Contents field into the document.

    The TOC field auto-collects headings of specified levels.
    User must right-click and 'Update Field' to populate.

    Args:
        doc: python-docx Document.
        heading_depth: Max heading level to include (1-9).
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)

    # TOC field code
    toc_code = f'TOC \\o "1-{heading_depth}" \\h \\z \\u'

    insert_field(p, toc_code)

    # Add instruction text
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(12)
    run = p2.add_run(
        "[To update: right-click the Table of Contents above "
        "and select 'Update Field', then 'Update entire table']"
    )
    run.font.size = Pt(8)
    run.font.name = config.FONT_NAME
    run.font.italic = True
    run.font.color.rgb = None  # default color


def add_styles_for_toc(doc):
    """Ensure Heading 1-3 styles are configured for TOC collection.

    Without proper heading styles, the TOC field will not find entries.
    This function sets up python-docx's built-in heading styles.
    """
    from docx.shared import Pt, RGBColor

    from tflshell import config

    F = config.FONT_NAME

    levels = [
        (doc.styles["Heading 1"], 14, config.HEADER_BG_HEX),
        (doc.styles["Heading 2"], 12, "1F4E79"),
        (doc.styles["Heading 3"], 11, "2E75B6"),
        (doc.styles["Heading 4"], 10, "000000"),
    ]

    for style, size, color_hex in levels:
        style.font.name = F
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color_hex)
        style.paragraph_format.keep_with_next = True
