"""Sponsor/Protocol/Page header block builder for TFL pages.

Every TFL page must carry this header block ABOVE the three-line table:

Sponsor: [Sponsor Name]                                 Page X of Y
Protocol: [Protocol Number]
[Study Title / Compound Name]

Table 14.X.X  [Full Descriptive Title]

Analysis Set: [ITT / Safety / PP Population]
"""

from docx.enum.text import WD_TAB_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor

from tflshell import config
from tflshell.docx_utils.page_numbering import add_page_number_fields
from tflshell.docx_utils.xml_helpers import insert_bookmark


def add_tfl_page_header(doc, tfl_item):
    """Add the complete TFL header block before a table/figure/listing.

    This includes sponsor/protocol info, page numbering, TFL ID,
    descriptive title, analysis population, and a bookmark for TOC linking.

    Args:
        doc: python-docx Document.
        tfl_item: TFLItem with header metadata.
    """
    F = config.FONT_NAME
    FS = config.FONT_SIZE_HEADER_BLOCK

    # Line 1: Sponsor + Page X of Y
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(f"Sponsor: {tfl_item.sponsor_placeholder}")
    run.font.size = Pt(FS)
    run.font.name = F

    content_width = config.PAGE_WIDTH_INCHES - (2 * config.MARGIN_INCHES)
    p.paragraph_format.tab_stops.add_tab_stop(
        Inches(content_width), alignment=WD_TAB_ALIGNMENT.RIGHT
    )
    p.add_run("\t")
    add_page_number_fields(p, prefix="Page ", suffix="")

    # Line 2: Protocol
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True

    run = p.add_run(f"Protocol: {tfl_item.protocol_placeholder}")
    run.font.size = Pt(FS)
    run.font.name = F

    # Line 3: Study Title
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(config.STUDY_TITLE_PLACEHOLDER)
    run.font.size = Pt(FS)
    run.font.name = F
    run.font.italic = True

    # Line 4: One visible black TFL heading. Heading 4 supplies the TOC target;
    # the bookmark is anchored here rather than in a duplicate italic heading.
    p = doc.add_paragraph(style=doc.styles["Heading 4"])
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(f"{tfl_item.display_label}  {tfl_item.title}")
    run.font.size = Pt(config.FONT_SIZE_H3)
    run.font.name = F
    run.font.bold = True
    run.font.italic = False
    run.font.color.rgb = RGBColor(0, 0, 0)
    insert_bookmark(p, tfl_item.id)

    # Badge
    badge = tfl_item.badge
    if badge:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(badge)
        run.font.size = Pt(config.FONT_SIZE_BADGE)
        run.font.name = F
        run.font.bold = True
        run.font.color.rgb = None  # Will use doc-level badge style

    # Line 6: Analysis Population
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(f"Analysis Set: {tfl_item.population}")
    run.font.size = Pt(FS)
    run.font.name = F
    run.font.italic = True
