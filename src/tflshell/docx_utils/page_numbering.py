"""Word PAGE and NUMPAGES field insertion utilities."""

from docx.shared import Pt
from tflshell import config
from tflshell.docx_utils.xml_helpers import insert_field


def add_page_number_fields(paragraph, prefix="Page ", suffix=""):
    """Insert 'Page X of Y' field codes into a paragraph.

    Uses Word's native PAGE and NUMPAGES fields which auto-update
    when the document is opened or printed.

    Args:
        paragraph: python-docx Paragraph object.
        prefix: Text before the page number.
        suffix: Text after the total pages.
    """
    run = paragraph.add_run(prefix)
    run.font.size = Pt(config.FONT_SIZE_HEADER_BLOCK)
    run.font.name = config.FONT_NAME

    insert_field(paragraph, "PAGE")

    run = paragraph.add_run(" of ")
    run.font.size = Pt(config.FONT_SIZE_HEADER_BLOCK)
    run.font.name = config.FONT_NAME

    insert_field(paragraph, "NUMPAGES")

    if suffix:
        run = paragraph.add_run(suffix)
        run.font.size = Pt(config.FONT_SIZE_HEADER_BLOCK)
        run.font.name = config.FONT_NAME
