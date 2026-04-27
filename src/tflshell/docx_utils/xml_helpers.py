"""Low-level OxmlElement utility functions for python-docx XML manipulation."""

from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def make_element(tag: str, attribs: dict = None) -> OxmlElement:
    """Create an OxmlElement with optional attributes.

    Args:
        tag: Element tag name in 'w:localName' format.
        attribs: Dict of attribute name -> value in 'w:attrName' format.
    """
    el = OxmlElement(tag)
    if attribs:
        for key, value in attribs.items():
            el.set(qn(key), str(value))
    return el


def make_field_begin() -> OxmlElement:
    """Create a w:fldChar with fldCharType='begin'."""
    return make_element("w:fldChar", {"w:fldCharType": "begin"})


def make_field_separate() -> OxmlElement:
    """Create a w:fldChar with fldCharType='separate'."""
    return make_element("w:fldChar", {"w:fldCharType": "separate"})


def make_field_end() -> OxmlElement:
    """Create a w:fldChar with fldCharType='end'."""
    return make_element("w:fldChar", {"w:fldCharType": "end"})


def make_instr_text(text: str) -> OxmlElement:
    """Create a w:instrText element with preserved whitespace."""
    el = OxmlElement("w:instrText")
    el.set(qn("xml:space"), "preserve")
    el.text = text
    return el


def make_border(edge: str, val: str = "single", sz: str = "4",
                space: str = "0", color: str = "000000") -> OxmlElement:
    """Create a w:border element for a specific edge.

    Args:
        edge: 'top', 'left', 'bottom', 'right', 'insideH', or 'insideV'.
        val: Border style ('single', 'none', 'double', etc.).
        sz: Border width in eighths of a point.
        space: Spacing offset.
        color: Hex color without '#'.
    """
    return make_element(f"w:{edge}", {
        "w:val": val,
        "w:sz": sz,
        "w:space": space,
        "w:color": color,
    })


def insert_field(paragraph, field_code: str):
    """Insert a Word field (e.g., PAGE, NUMPAGES, TOC) into a paragraph.

    The field is inserted at the current end of the paragraph using a run.
    """
    run = paragraph.add_run()
    run._r.append(make_field_begin())
    run2 = paragraph.add_run()
    run2._r.append(make_instr_text(field_code))
    run3 = paragraph.add_run()
    run3._r.append(make_field_separate())
    run4 = paragraph.add_run()
    run4._r.append(make_field_end())


def insert_bookmark(paragraph, bookmark_name: str, bookmark_id: str = "0"):
    """Insert a Word bookmark around a paragraph's content.

    The bookmark tags are inserted via a new run at the end of the paragraph.
    """
    run = paragraph.add_run()
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bookmark_id)
    start.set(qn("w:name"), bookmark_name)
    run._r.append(start)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bookmark_id)
    run._r.append(end)


def set_table_borders(table, top_sz=12, bottom_sz=12, thin_sz=4,
                      color="000000"):
    """Apply three-line table borders at the table level.

    Sets: thick top, thick bottom, none for left/right/insideH/insideV.
    Returns the w:tblBorders element.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)

    borders = OxmlElement("w:tblBorders")

    borders.append(make_border("top", sz=str(top_sz), color=color))
    borders.append(make_border("bottom", sz=str(bottom_sz), color=color))
    borders.append(make_border("left", val="none", sz="0", color="auto"))
    borders.append(make_border("right", val="none", sz="0", color="auto"))
    borders.append(make_border("insideH", val="none", sz="0", color="auto"))
    borders.append(make_border("insideV", val="none", sz="0", color="auto"))

    tblPr.append(borders)
    return borders


def set_cell_bottom_border(cell, sz=4, color="000000"):
    """Apply a thin bottom border to an individual table cell (used for header row)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:tcBorders"))
    if existing is not None:
        tcPr.remove(existing)

    borders = OxmlElement("w:tcBorders")
    borders.append(make_border("bottom", sz=str(sz), color=color))
    tcPr.append(borders)


def set_cell_text(cell, text: str, bold=False, font_size=9,
                  font_name="Times New Roman", alignment=None,
                  space_before=1, space_after=1):
    """Clear and set text in a table cell with consistent formatting.

    Args:
        space_before: Paragraph space before in points. Use 4-6 for category headers.
        space_after: Paragraph space after in points.
    """
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Clear existing paragraphs
    for p in cell.paragraphs:
        p.clear()

    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(12)

    run = p.add_run(str(text))
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    return p


def set_cell_shading(cell, color_hex: str):
    """Apply background shading to a table cell.

    Args:
        cell: python-docx cell object.
        color_hex: Background color in hex format (e.g. 'F2F2F2') without '#'.
    """
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_width(cell, width_cm: float):
    """Set explicit width on a table cell (controls column width).

    Args:
        cell: python-docx cell object.
        width_cm: Width in centimeters.
    """
    from docx.shared import Cm
    cell.width = Cm(width_cm)
