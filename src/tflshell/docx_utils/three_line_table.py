"""Three-line table (三线表) format for python-docx tables v2.2.

Supports:
  - Rich row metadata (bold, indent) via dict-format data_rows
  - Gray header background for visual distinction
  - Left-aligned parameter column, center-aligned data columns
  - Explicit column widths (param col wider than data cols)
  - Backward-compatible with flat list[str] data_rows
"""

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from tflshell import config
from tflshell.docx_utils.xml_helpers import (
    set_cell_bottom_border,
    set_cell_shading,
    set_cell_text,
    set_table_autofit_to_page,
    set_table_borders,
)
from tflshell.presentation import get_presentation_profile


def apply_three_line_format(table, thick_sz=None, thin_sz=None, color=None):
    thick_sz = thick_sz or config.THREE_LINE_THICK_SZ
    thin_sz = thin_sz or config.THREE_LINE_THIN_SZ
    color = color or config.THREE_LINE_BORDER_COLOR
    set_table_borders(table, top_sz=thick_sz, bottom_sz=thick_sz, thin_sz=thin_sz, color=color)
    if table.rows:
        for cell in table.rows[0].cells:
            set_cell_bottom_border(cell, sz=thin_sz, color=color)


def _split_header_text(text: str) -> list[str]:
    """Split a column header into multiple lines for readability."""
    if "\n" in text:
        return text.split("\n")

    if " / " in text and len(text) > 18:
        parts = text.split(" / ", 1)
        return [parts[0] + " /", parts[1]]

    if "(N=" in text and len(text) > 12:
        idx = text.index("(N=")
        if idx > 0:
            return [text[:idx].strip(), text[idx:].strip()]

    return [text]


def _format_header_cell(cell, header_text: str, alignment=None):
    """Set multi-line header text in a cell, bold."""
    if alignment is None:
        alignment = WD_ALIGN_PARAGRAPH.LEFT
    lines = _split_header_text(header_text)

    for p in cell.paragraphs:
        p.clear()

    if len(lines) == 1:
        p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        p.alignment = alignment
        run = p.add_run(lines[0])
        run.font.name = config.FONT_NAME
        run.font.size = Pt(config.FONT_SIZE_TABLE)
        run.font.bold = True
    else:
        first = True
        for line in lines:
            if first:
                p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                first = False
            else:
                p = cell.add_paragraph()
            p.alignment = alignment
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(line)
            run.font.name = config.FONT_NAME
            run.font.size = Pt(config.FONT_SIZE_TABLE)
            run.font.bold = True


def _alignment(value, default=WD_ALIGN_PARAGRAPH.CENTER):
    """Resolve a declarative alignment name to python-docx."""
    if value == "left":
        return WD_ALIGN_PARAGRAPH.LEFT
    if value == "right":
        return WD_ALIGN_PARAGRAPH.RIGHT
    return default


def _set_repeat_table_header(row):
    """Mark a header row to repeat after an automatic page break."""
    tr_pr = row._tr.get_or_add_trPr()
    repeat = tr_pr.find(qn("w:tblHeader"))
    if repeat is None:
        repeat = OxmlElement("w:tblHeader")
        tr_pr.append(repeat)
    repeat.set(qn("w:val"), "true")


def _render_header_rows(table, header_rows, leaf_cols):
    """Render declarative header rows with horizontal/vertical spans."""
    occupied = [[False] * leaf_cols for _ in header_rows]
    for row_idx, header_row in enumerate(header_rows):
        col_idx = 0
        for spec in header_row:
            while col_idx < leaf_cols and occupied[row_idx][col_idx]:
                col_idx += 1
            colspan = max(int(spec.get("colspan", 1)), 1)
            rowspan = max(int(spec.get("rowspan", 1)), 1)
            if col_idx + colspan > leaf_cols or row_idx + rowspan > len(header_rows):
                raise ValueError(f"Header span exceeds declared grid at row {row_idx + 1}")
            start = table.cell(row_idx, col_idx)
            end = table.cell(row_idx + rowspan - 1, col_idx + colspan - 1)
            cell = start.merge(end) if (colspan > 1 or rowspan > 1) else start
            _format_header_cell(
                cell,
                str(spec.get("label", "")),
                alignment=_alignment(
                    spec.get("alignment"),
                    WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER,
                ),
            )
            set_cell_shading(cell, config.HEADER_ROW_BG_HEX)
            for rr in range(row_idx, row_idx + rowspan):
                for cc in range(col_idx, col_idx + colspan):
                    occupied[rr][cc] = True
            col_idx += colspan
        if not all(occupied[row_idx]):
            raise ValueError(f"Header row {row_idx + 1} does not cover {leaf_cols} columns")
        _set_repeat_table_header(table.rows[row_idx])


def _render_data_row(
    table,
    row_idx,
    row_data,
    col_count,
    font_size,
    font_name,
    profile,
    column_alignments=None,
):
    """Render a single data row with rich formatting support.

    Spacing rules:
      - Bold category headers: extra space above (6pt) to separate from previous group
      - Indented sub-items: tight spacing (0pt above, 1pt after) within the group
      - Normal rows: standard 1pt spacing
    """
    if isinstance(row_data, dict):
        label = row_data.get("label", "")
        bold = row_data.get("bold", False)
        indent = row_data.get("indent", False)
        values = row_data.get("values", [])
    else:
        label = row_data[0] if len(row_data) > 0 else ""
        values = list(row_data[1:]) if len(row_data) > 1 else []
        bold = False
        indent = False

    expected_value_count = max(col_count - 1, 0)
    if len(values) != expected_value_count:
        raise ValueError(
            f"Row '{label}' has {len(values)} value cells but expected "
            f"{expected_value_count} based on the table header."
        )

    # Spacing: bold category headers get extra space above to separate groups
    if bold:
        space_before = profile.table.group_header_space_before
        space_after = profile.table.group_header_space_after
    elif indent:
        space_before = profile.table.indented_row_space_before
        space_after = profile.table.indented_row_space_after
    else:
        space_before = profile.table.standard_row_space_before
        space_after = profile.table.standard_row_space_after

    # Column 0: Parameter / Label — left-aligned, optional bold/indent
    display_label = (profile.table.indent_prefix if indent else "") + label
    set_cell_text(
        table.cell(row_idx, 0),
        display_label,
        bold=bold,
        font_size=font_size,
        font_name=font_name,
        alignment=_alignment(
            column_alignments[0] if column_alignments else "left", WD_ALIGN_PARAGRAPH.LEFT
        ),
        space_before=space_before,
        space_after=space_after,
        line_spacing=profile.table.cell_line_spacing,
    )

    # Columns 1+: Numeric values — center-aligned, inherit bold from row
    for col_idx in range(1, col_count):
        val = values[col_idx - 1]
        set_cell_text(
            table.cell(row_idx, col_idx),
            val,
            bold=bold,
            font_size=font_size,
            font_name=font_name,
            alignment=_alignment(
                column_alignments[col_idx]
                if column_alignments and col_idx < len(column_alignments)
                else "center"
            ),
            space_before=space_before,
            space_after=space_after,
            line_spacing=profile.table.cell_line_spacing,
        )


def create_three_line_table(
    doc,
    rows,
    cols,
    headers,
    data_rows=None,
    font_size=None,
    font_name=None,
    presentation_profile="csr_standard",
    header_rows=None,
    column_alignments=None,
):
    """Create a three-line table with rich formatting support.

    Args:
        doc: python-docx Document.
        rows: Total number of rows (including header).
        cols: Number of columns.
        headers: List of column header strings.
        data_rows: List of dicts or list[list[str]] for data rows.
        font_size: Override default font size.
        font_name: Override default font name.

    Returns:
        python-docx Table object.
    """
    font_size = font_size or config.FONT_SIZE_TABLE
    font_name = font_name or config.FONT_NAME
    profile = get_presentation_profile(presentation_profile)

    header_rows = header_rows or [[{"label": text} for text in headers]]
    header_row_count = len(header_rows)
    data_row_count = max(rows - 1, 0)
    table = doc.add_table(rows=header_row_count + data_row_count, cols=cols)
    table.autofit = True
    set_table_autofit_to_page(table, width_pct=profile.table.width_pct)

    # ---- Header rows ----
    if len(header_rows) == 1 and all(
        "colspan" not in x and "rowspan" not in x for x in header_rows[0]
    ):
        for col_idx, spec in enumerate(header_rows[0]):
            cell = table.rows[0].cells[col_idx]
            align = _alignment(
                spec.get("alignment"),
                WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER,
            )
            _format_header_cell(cell, str(spec.get("label", "")), alignment=align)
            set_cell_shading(cell, config.HEADER_ROW_BG_HEX)
        _set_repeat_table_header(table.rows[0])
    else:
        _render_header_rows(table, header_rows, cols)

    # ---- Data rows ----
    if data_rows:
        for data_idx in range(data_row_count):
            row_idx = header_row_count + data_idx
            if data_idx < len(data_rows):
                _render_data_row(
                    table,
                    row_idx,
                    data_rows[data_idx],
                    cols,
                    font_size,
                    font_name,
                    profile,
                    column_alignments,
                )
            else:
                # Fill extra rows with XX
                for col_idx in range(cols):
                    set_cell_text(
                        table.cell(row_idx, col_idx),
                        "XX",
                        font_size=font_size,
                        font_name=font_name,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER,
                        line_spacing=profile.table.cell_line_spacing,
                    )
    else:
        # No data rows provided: fill with XX placeholders
        for row_idx in range(header_row_count, header_row_count + data_row_count):
            for col_idx in range(cols):
                align = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
                set_cell_text(
                    table.cell(row_idx, col_idx),
                    "XX",
                    font_size=font_size,
                    font_name=font_name,
                    alignment=align,
                    line_spacing=profile.table.cell_line_spacing,
                )

    # ---- Apply three-line borders ----
    apply_three_line_format(table)
    for cell in table.rows[header_row_count - 1].cells:
        set_cell_bottom_border(
            cell, sz=config.THREE_LINE_THIN_SZ, color=config.THREE_LINE_BORDER_COLOR
        )
    return table
