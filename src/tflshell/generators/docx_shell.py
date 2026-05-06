"""DOCX Shell Template Generator.

Produces the TFL Shell Template with:
- Three-line table format (三线表)
- Sponsor/Protocol/Page header blocks on every TFL
- Embedded matplotlib-generated clinical figures
- Word-native Table of Contents with bookmarks
- Dataset lineage and traceability footnotes
"""

import os
from datetime import datetime

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from tflshell import config
from tflshell.docx_utils.header_block import add_tfl_page_header
from tflshell.docx_utils.three_line_table import create_three_line_table
from tflshell.docx_utils.toc_builder import add_styles_for_toc, insert_toc_field
from tflshell.models.catalog import TFLCatalog
from tflshell.models.enums import Section, TFLType
from tflshell.presentation import get_presentation_profile
from tflshell.utils import styles as sty
from tflshell.utils.naming import make_filename


class DocxShellGenerator:
    """Generate the TFL Shell Template DOCX document."""

    def __init__(
        self,
        catalog: TFLCatalog,
        output_path: str | None = None,
        therapeutic_area: str = "all",
        generate_figures: bool = True,
        sponsor: str = None,
        protocol: str = None,
        presentation_profile: str = "csr_standard",
    ):
        self.catalog = catalog
        self.output_path = output_path or os.path.join(
            config.DEFAULT_OUTPUT_DIR,
            make_filename("TFL_Shell_Template", config.VERSION, ".docx"),
        )
        self.therapeutic_area = therapeutic_area
        self.generate_figures = generate_figures
        self.sponsor_override = sponsor
        self.protocol_override = protocol
        self.presentation_profile = get_presentation_profile(presentation_profile)
        self.doc: Document = None

    def generate(self) -> str:
        self.doc = Document()
        self._setup_document()
        add_styles_for_toc(self.doc)
        self._build_cover_page()
        self._build_toc()
        self._build_introduction()
        self._build_tfl_sections()
        self._set_document_properties()
        os.makedirs(os.path.dirname(self.output_path), exist_ok=True)
        self.doc.save(self.output_path)
        return self.output_path

    def _setup_document(self):
        section = self.doc.sections[0]
        section.page_width = Inches(config.PAGE_WIDTH_INCHES)
        section.page_height = Inches(config.PAGE_HEIGHT_INCHES)
        section.orientation = WD_ORIENT.LANDSCAPE
        section.top_margin = Inches(config.MARGIN_INCHES)
        section.bottom_margin = Inches(config.MARGIN_INCHES)
        section.left_margin = Inches(config.MARGIN_INCHES)
        section.right_margin = Inches(config.MARGIN_INCHES)
        sty.set_default_font(self.doc, config.FONT_NAME, config.FONT_SIZE_BODY)

        # Header
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = hp.add_run(f"[Protocol Number] — TFL Shell Template v{config.VERSION}")
        run.font.size = Pt(8)
        run.font.name = config.FONT_NAME

        # Footer
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        run = fp.add_run("CONFIDENTIAL")
        run.font.size = Pt(8)
        run.font.name = config.FONT_NAME

    def _build_cover_page(self):
        doc = self.doc
        for _ in range(5):
            doc.add_paragraph("")

        p = doc.add_paragraph("")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Clinical Study Report\nTFL Shell Template")
        run.font.size = Pt(config.FONT_SIZE_COVER_TITLE)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(config.HEADER_BG_HEX)
        run.font.name = config.FONT_NAME

        doc.add_paragraph("")

        subtitle = doc.add_paragraph("")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(
            "Tables, Figures, and Listings\n"
            "ICH E3 Sections 14.1 / 14.2 / 14.3 / 14.4 / 16.2\n\n"
            "Automatic Word TOC | Three-Line Table Format (三线表)\n"
            "Embedded Clinical Figures | Dataset Traceability"
        )
        run.font.size = Pt(12)
        run.font.name = config.FONT_NAME

        doc.add_paragraph("")

        for line in [
            f"Version: {config.VERSION}",
            f"Date: {datetime.now().strftime('%d %B %Y')}",
            "Classification: CONFIDENTIAL",
            "",
            "Applicable to: Oncology and Non-Oncology Clinical Studies",
            "Template Type: Master Shell — No Patient Data",
        ]:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(11)
                run.font.name = config.FONT_NAME

        doc.add_page_break()

    def _build_toc(self):
        """Insert a Word-native Table of Contents field."""
        doc = self.doc
        doc.add_paragraph("Table of Contents", style=doc.styles["Heading 1"])
        insert_toc_field(doc, heading_depth=4)
        doc.add_page_break()

    def _build_introduction(self):
        doc = self.doc
        doc.add_paragraph("1.0  Introduction and Usage Notes", style=doc.styles["Heading 1"])
        doc.add_paragraph(
            "This document is the master TFL (Tables, Figures, and Listings) shell "
            "template for clinical study reports covering CSR Sections 14.1, 14.2, 14.3, "
            "14.4, and 16.2 only; Section 16.1 is intentionally excluded from this generator. "
            "Each shell includes a sponsor/protocol header block, "
            "descriptive title, analysis population, shell-first table/listing structure or "
            "simulated figure, and traceability footnotes. No actual study results are "
            "included in this template."
        )

        doc.add_paragraph("1.1  Table and Listing Shell Conventions", style=doc.styles["Heading 2"])
        doc.add_paragraph(
            "All tables and listings use shell-first structure. The first column preserves "
            "the example row labels, categories, or subject-structure identifiers needed to "
            "show the intended layout. Non-structural cells use style-appropriate shell "
            "placeholders such as XX, xx (xx.x), x.xxx, or CI-style formats rather than "
            "mock numeric or subject-level results. Controlled tables use `Group 1`, `Group 2`, "
            "and an optional separate ellipsis (`...`) expansion column; that expansion column "
            "must never be merged with Overall, HR, Total, or other analytic columns. Tables "
            "use the standard three-line table format (三线表) and should auto-fit within the page."
        )

        doc.add_paragraph("1.2  How to Use This Template", style=doc.styles["Heading 2"])
        doc.add_paragraph(
            "1. Select applicable TFLs using the companion XLSX catalog workbook.\n"
            "2. Replace bracketed metadata placeholders such as sponsor and protocol details.\n"
            "3. Keep general shells plus only the oncology-only or non-oncology-only shells that apply to the study.\n"
            "4. Verify titles, headers, placeholders, and footnotes against the protocol, SAP, and shell standards.\n"
            "5. In Word, update the Table of Contents field to populate headings.\n"
            "6. Use the shells as governance-ready templates for downstream programming and review."
        )

        doc.add_paragraph("1.3  Figure Shells", style=doc.styles["Heading 2"])
        doc.add_paragraph(
            "Figures retain simulated illustrations generated by matplotlib so the output "
            "shows expected visual layout for Kaplan-Meier curves, waterfall plots, spider "
            "plots, swimmer plots, forest plots, box plots, and longitudinal plots. These "
            "images are for shell illustration only and must be replaced with study-specific "
            "outputs in production. Page layout is optimized to keep each figure shell on one "
            "page when reasonably possible, but this remains a best-effort rule."
        )

        doc.add_page_break()

    def _build_tfl_sections(self):
        sections = [
            ("14.1  Demographics and Baseline Characteristics", Section.SEC_14_1),
            ("14.2  Efficacy Analysis", Section.SEC_14_2),
            ("14.3  Safety Analysis", Section.SEC_14_3),
            ("14.4  Special Assessments", Section.SEC_14_4),
            ("16.2  Patient Data Listings", Section.SEC_16_2),
        ]

        for heading, section_enum in sections:
            self.doc.add_paragraph(heading, style=self.doc.styles["Heading 2"])

            items = self.catalog.by_section(section_enum)
            if self.therapeutic_area != "all":
                area = self._normalize_area(self.therapeutic_area)
                items = [i for i in items if area in i.therapeutic_areas]

            # Sort: tables first, then figures, then listings
            tables = [i for i in items if i.tfl_type == TFLType.TABLE]
            figures = [i for i in items if i.tfl_type == TFLType.FIGURE]
            listings = [i for i in items if i.tfl_type == TFLType.LISTING]
            tables.sort(key=lambda i: i.sort_key)
            figures.sort(key=lambda i: i.sort_key)
            listings.sort(key=lambda i: i.sort_key)

            # For 14.3, group by sub-section
            if section_enum == Section.SEC_14_3:
                self._render_143_sub_sections(tables, figures, listings)
                continue

            if tables:
                self._add_sub_heading("Tables")
                for item in tables:
                    self._add_tfl_shell(item)

            if figures:
                self._add_sub_heading("Figures")
                for item in figures:
                    self._add_tfl_shell(item)

            if listings:
                self._add_sub_heading("Listings")
                for item in listings:
                    self._add_tfl_shell(item)

    def _render_143_sub_sections(self, tables, figures, listings):
        """Render 14.3 with sub-section headings (14.3.1 AE, 14.3.2 Other, etc.)."""
        sub_sections = [
            ("14.3.1  Adverse Events", "1"),
            ("14.3.2  Other Safety Observations", "2"),
            ("14.3.3  Clinical Laboratory Evaluations", "3"),
            ("14.3.4  Vital Signs, ECG, and Physical Examinations", "4"),
        ]
        for sub_heading, sub_num in sub_sections:
            # Filter items by sub-number (first digit after section)
            sub_tables = [i for i in tables if i.id.split(".")[-2] == sub_num]
            sub_figs = [i for i in figures if i.id.split(".")[-2] == sub_num]
            if not sub_tables and not sub_figs:
                continue
            self._add_sub_heading(sub_heading)
            if sub_tables:
                for item in sub_tables:
                    self._add_tfl_shell(item)
            if sub_figs:
                for item in sub_figs:
                    self._add_tfl_shell(item)

    def _add_sub_heading(self, text):
        """Add a sub-heading for 'Tables' / 'Figures' / 'Listings' groups."""
        p = self.doc.add_paragraph(style=self.doc.styles["Heading 3"])
        p.paragraph_format.space_before = Pt(
            self.presentation_profile.paragraphs.subheading_space_before
        )
        p.paragraph_format.space_after = Pt(
            self.presentation_profile.paragraphs.subheading_space_after
        )
        run = p.add_run(text)
        run.font.name = config.FONT_NAME
        run.font.size = Pt(config.FONT_SIZE_H3)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(config.HEADER_BG_HEX)

    def _add_tfl_shell_heading(self, tfl):
        """TFL title as Heading 4 — appears in TOC under sub-group."""
        p = self.doc.add_paragraph(style=self.doc.styles["Heading 4"])
        p.paragraph_format.space_before = Pt(
            self.presentation_profile.paragraphs.shell_heading_space_before
        )
        p.paragraph_format.space_after = Pt(
            self.presentation_profile.paragraphs.shell_heading_space_after
        )
        p.paragraph_format.keep_with_next = True
        run = p.add_run(f"{tfl.display_label}  {tfl.title}")
        run.font.name = config.FONT_NAME
        run.font.size = Pt(10)
        run.font.bold = True

    def _add_tfl_shell(self, tfl):
        doc = self.doc

        # Override sponsor/protocol if provided
        if self.sponsor_override:
            tfl.sponsor_placeholder = self.sponsor_override
        if self.protocol_override:
            tfl.protocol_placeholder = self.protocol_override

        self._add_tfl_shell_heading(tfl)

        # Header block (sponsor, protocol, page, title, population) + bookmark
        add_tfl_page_header(doc, tfl)

        # Table / Figure / Listing body
        if tfl.tfl_type == TFLType.TABLE:
            self._add_table_body(tfl)
        elif tfl.tfl_type == TFLType.FIGURE:
            self._add_figure_body(tfl)
        elif tfl.tfl_type == TFLType.LISTING:
            self._add_listing_body(tfl)

        # Footnotes
        all_notes = tfl.footnote_text()
        if all_notes:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(
                self.presentation_profile.paragraphs.footnote_label_space_before
            )
            p.paragraph_format.space_after = Pt(
                self.presentation_profile.paragraphs.footnote_label_space_after
            )
            run = p.add_run("Footnotes:")
            run.font.size = Pt(config.FONT_SIZE_FOOTNOTE)
            run.font.bold = True
            run.font.italic = True
            run.font.name = config.FONT_NAME
            p.paragraph_format.keep_with_next = True

            for i, note in enumerate(all_notes, 1):
                p = self.doc.add_paragraph()
                p.paragraph_format.space_before = Pt(
                    self.presentation_profile.paragraphs.footnote_item_space_before
                )
                p.paragraph_format.space_after = Pt(
                    self.presentation_profile.paragraphs.footnote_item_space_after
                )
                run = p.add_run(f"[{i}] {note}")
                run.font.size = Pt(config.FONT_SIZE_FOOTNOTE)
                run.font.italic = True
                run.font.name = config.FONT_NAME

        # Page break after each TFL
        doc.add_page_break()

    def _add_table_body(self, tfl):
        cols = tfl.placeholder_columns or ["Column 1", "Column 2", "Column 3"]
        data_rows = tfl.shell_data_rows_rich
        total_rows = len(data_rows) + 1  # +1 for header

        create_three_line_table(
            self.doc,
            rows=total_rows,
            cols=len(cols),
            headers=cols,
            data_rows=data_rows,
            presentation_profile=self.presentation_profile.name,
        )

    def _add_figure_body(self, tfl):
        doc = self.doc

        if self.generate_figures and tfl.is_figure_generated:
            try:
                from tflshell.generators.figure_engine import generate_figure_buffer

                buf = generate_figure_buffer(tfl)
                # Insert image directly
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_before = Pt(
                    self.presentation_profile.paragraphs.figure_space_before
                )
                p_img.paragraph_format.space_after = Pt(
                    self.presentation_profile.paragraphs.figure_space_after
                )
                run_img = p_img.add_run()
                run_img.add_picture(buf, width=Inches(min(tfl.figure_width_inches, 6.0)))
                # Figure description as italic caption
                if tfl.figure_description:
                    p_desc = doc.add_paragraph()
                    p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_desc.paragraph_format.space_before = Pt(
                        self.presentation_profile.paragraphs.caption_space_before
                    )
                    p_desc.paragraph_format.space_after = Pt(
                        self.presentation_profile.paragraphs.caption_space_after
                    )
                    run_d = p_desc.add_run(f"[{tfl.display_label}: {tfl.figure_description}]")
                    run_d.font.size = Pt(8)
                    run_d.font.italic = True
                    run_d.font.name = config.FONT_NAME
                    run_d.font.color.rgb = RGBColor(100, 100, 100)
                # Gray shell note
                self._add_figure_shell_note()
                return
            except Exception as e:
                print(f"  WARNING: Figure generation failed for {tfl.id}: {e}")

        # Fallback: centered placeholder text
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(30)
        p.paragraph_format.space_after = Pt(10)
        run = p.add_run("INSERT FIGURE HERE")
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(128, 128, 128)
        run.font.name = config.FONT_NAME
        if tfl.figure_description:
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = p2.add_run(tfl.figure_description)
            run2.font.size = Pt(10)
            run2.font.italic = True
            run2.font.name = config.FONT_NAME
        self._add_figure_shell_note()

    def _add_figure_shell_note(self):
        """Add gray italic shell template note below figure."""
        p_shell = self.doc.add_paragraph()
        p_shell.paragraph_format.space_before = Pt(
            self.presentation_profile.paragraphs.figure_note_space_before
        )
        p_shell.paragraph_format.space_after = Pt(
            self.presentation_profile.paragraphs.figure_note_space_after
        )
        run = p_shell.add_run(
            "[Figure shell template. Actual figure generated by clinical statistical software (e.g., SAS, R).]"
        )
        run.font.size = Pt(config.FONT_SIZE_FOOTNOTE)
        run.font.color.rgb = RGBColor(*config.FIGURE_SHELL_NOTE_COLOR)
        run.font.italic = True
        run.font.name = config.FONT_NAME

    def _add_listing_body(self, tfl):
        doc = self.doc
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(
            self.presentation_profile.paragraphs.listing_note_space_before
        )
        p.paragraph_format.space_after = Pt(
            self.presentation_profile.paragraphs.listing_note_space_after
        )
        p.paragraph_format.keep_with_next = True
        run = p.add_run("Note: Sorted by site/subject ID.")
        run.font.size = Pt(config.FONT_SIZE_FOOTNOTE)
        run.font.italic = True
        run.font.name = config.FONT_NAME

        if tfl.table_notes:
            run2 = p.add_run(f" {tfl.table_notes}")
            run2.font.size = Pt(config.FONT_SIZE_FOOTNOTE)
            run2.font.name = config.FONT_NAME

        cols = tfl.placeholder_columns or ["Variable 1", "Variable 2", "Variable 3"]
        data_rows = tfl.shell_data_rows
        total_rows = len(data_rows) + 1

        create_three_line_table(
            self.doc,
            rows=total_rows,
            cols=len(cols),
            headers=cols,
            data_rows=data_rows,
            font_size=config.FONT_SIZE_TABLE - 1,
            presentation_profile=self.presentation_profile.name,
        )

    def _set_document_properties(self):
        doc = self.doc
        doc.core_properties.title = f"Clinical Study Report — TFL Shell Template v{config.VERSION}"
        doc.core_properties.subject = "Tables, Figures, and Listings (Three-Line Format)"
        doc.core_properties.version = config.VERSION
        doc.core_properties.category = "Clinical Statistics"

    @staticmethod
    def _normalize_area(raw: str) -> str:
        cleaned = raw.strip().lower().replace("_", "-").replace(" ", "-")
        if cleaned == "oncology":
            return "Oncology"
        if cleaned in ("non-oncology", "nononcology"):
            return "Non-Oncology"
        return cleaned
