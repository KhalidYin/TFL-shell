"""SOP DOCX Generator v2.1."""

import os

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from tflshell import config
from tflshell.data.sop_content import build_sop_content
from tflshell.docx_utils.toc_builder import add_styles_for_toc, insert_toc_field
from tflshell.utils.naming import make_filename


class DocxSopGenerator:
    """Generate the TFL Shell SOP document."""

    def __init__(self, output_path: str | None = None):
        self.output_path = output_path or os.path.join(
            config.DEFAULT_OUTPUT_DIR,
            make_filename("TFL_Shell_SOP", config.VERSION, ".docx"),
        )
        self.doc: Document = None

    def generate(self) -> str:
        self.doc = Document()
        self._setup_document()
        add_styles_for_toc(self.doc)
        self._build_content()
        self._set_document_properties()
        os.makedirs(os.path.dirname(self.output_path), exist_ok=True)
        self.doc.save(self.output_path)
        return self.output_path

    def _setup_document(self):
        section = self.doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)
        section.orientation = WD_ORIENT.PORTRAIT
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    def _build_content(self):
        sop = build_sop_content(config.VERSION)
        doc = self.doc

        self._add_header_block(sop)
        doc.add_page_break()

        toc_heading = doc.add_paragraph("Table of Contents", style=doc.styles["Heading 1"])
        for run in toc_heading.runs:
            run.font.name = config.FONT_NAME
            run.font.size = Pt(14)
            run.font.bold = True
        insert_toc_field(doc, heading_depth=3)
        doc.add_page_break()

        for section in sop.sections:
            h1 = doc.add_paragraph(
                f"{section['number']}.  {section['title']}", style=doc.styles["Heading 1"]
            )
            for run in h1.runs:
                run.font.name = config.FONT_NAME
                run.font.size = Pt(14)
                run.font.bold = True

            for sub in section["subsections"]:
                h2 = doc.add_paragraph(
                    f"{sub['number']}  {sub['title']}", style=doc.styles["Heading 2"]
                )
                for run in h2.runs:
                    run.font.name = config.FONT_NAME
                    run.font.size = Pt(12)
                    run.font.bold = True

                for para_text in sub["content"]:
                    p = doc.add_paragraph(para_text)
                    for run in p.runs:
                        run.font.name = config.FONT_NAME
                        run.font.size = Pt(11)

            doc.add_paragraph("")

    def _add_header_block(self, sop):
        doc = self.doc
        doc.add_paragraph("")
        p = doc.add_paragraph("STANDARD OPERATING PROCEDURE")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.name = config.FONT_NAME

        doc.add_paragraph("")
        p = doc.add_paragraph(f"Title: {sop.title}")
        for run in p.runs:
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.name = config.FONT_NAME

        doc.add_paragraph("")
        table = doc.add_table(rows=5, cols=2)
        table.style = "Light Grid Accent 1"
        header_data = [
            ("SOP No.", sop.sop_number),
            ("Version", sop.version),
            ("Effective Date", sop.effective_date),
            ("Department", sop.department),
            ("Classification", "CONFIDENTIAL"),
        ]
        for row_idx, (label, value) in enumerate(header_data):
            run = table.rows[row_idx].cells[0].paragraphs[0].add_run(label)
            run.font.name = config.FONT_NAME
            run.font.size = Pt(11)
            run.font.bold = True
            run = table.rows[row_idx].cells[1].paragraphs[0].add_run(value)
            run.font.name = config.FONT_NAME
            run.font.size = Pt(11)
            table.rows[row_idx].cells[0].width = Inches(2.5)
            table.rows[row_idx].cells[1].width = Inches(3.5)

    def _set_document_properties(self):
        self.doc.core_properties.title = "SOP: TFL Shell Template Usage and Management"
        self.doc.core_properties.subject = "Standard Operating Procedure"
        self.doc.core_properties.version = config.SOP_VERSION
        self.doc.core_properties.category = "SOP"
