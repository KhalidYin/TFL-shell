from docx import Document
from openpyxl import load_workbook

from tflshell.data.definitions import build_catalog
from tflshell.generators.docx_shell import DocxShellGenerator
from tflshell.generators.docx_sop import DocxSopGenerator
from tflshell.generators.xlsx_toc import XlsxTocGenerator


REPRESENTATIVE_IDS = [
    "T14.2.26",
    "T14.2.28",
    "T14.2.30",
    "L16.2.35",
    "L16.2.36",
    "L16.2.37",
]

REMOVED_IDS = [
    "T14.2.3",
    "T14.2.17",
    "T14.2.18",
    "T14.3.1.11",
    "T14.3.1.12",
]

REMOVED_HEADINGS = [
    "Table 14.2.3  Subgroup Analysis of Primary Endpoint",
    "Table 14.2.17  PFS by Subgroup",
    "Table 14.2.18  OS by Subgroup",
    "Table 14.3.1.11  TEAEs by Age Group",
    "Table 14.3.1.12  TEAEs by Sex",
]


def _non_empty_paragraphs(doc: Document) -> list[str]:
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


def test_docx_and_workbook_stay_aligned_on_catalog_size_and_representative_ids(tmp_path):
    catalog = build_catalog()
    docx_path = tmp_path / "shell.docx"
    xlsx_path = tmp_path / "toc.xlsx"

    DocxShellGenerator(catalog, output_path=str(docx_path), generate_figures=False).generate()
    XlsxTocGenerator(catalog, output_path=str(xlsx_path)).generate()

    doc = Document(docx_path)
    doc_headings = [p.text for p in doc.paragraphs if p.style and p.style.name == "Heading 4"]

    workbook = load_workbook(xlsx_path, read_only=True, data_only=True)
    master_sheet = workbook["TOC_Master"]
    workbook_ids = [row[0] for row in master_sheet.iter_rows(min_row=2, values_only=True)]

    assert len(doc_headings) == len(workbook_ids) == len(catalog.all())

    for item_id in REPRESENTATIVE_IDS:
        item = catalog.get(item_id)
        assert item is not None
        assert item_id in workbook_ids
        assert f"{item.display_label}  {item.title}" in doc_headings


def test_sop_and_workbook_share_governance_language(tmp_path):
    xlsx_path = tmp_path / "toc.xlsx"
    sop_path = tmp_path / "sop.docx"

    XlsxTocGenerator(build_catalog(), output_path=str(xlsx_path)).generate()
    DocxSopGenerator(output_path=str(sop_path)).generate()

    workbook = load_workbook(xlsx_path, read_only=True, data_only=True)
    usage_sheet = workbook["Usage_Guide"]
    field_sheet = workbook["Field_Definitions"]
    usage_rows = {row[0]: row[1] for row in usage_sheet.iter_rows(min_row=2, values_only=True) if row[0]}
    field_names = {row[0] for row in field_sheet.iter_rows(min_row=2, values_only=True) if row[0]}

    sop_doc = Document(sop_path)
    sop_text = "\n".join(_non_empty_paragraphs(sop_doc))

    assert "14.1, 14.2, 14.3, 14.4, and 16.2" in usage_rows["Scope"]
    assert "14.1, 14.2, 14.3, 14.4, and 16.2" in sop_text
    assert "Coverage metadata" in usage_rows
    assert {"Shell Family", "Study Phase Scope", "Coverage Summary"} <= field_names
    assert "Group 1, Group 2" in usage_rows["Placeholder convention"]
    assert "ellipsis (...)" in usage_rows["Placeholder convention"]
    assert "must not be merged with Overall, Total, HR" in usage_rows["Placeholder convention"]
    assert "shell family, study phase scope, coverage summary" in sop_text
    assert "Automated quality gates should verify generation, catalog validation, and regression tests" in sop_text
    assert "ellipsis (`...`) expansion column" in sop_text


def test_removed_subgroup_tables_do_not_appear_in_docx_or_workbook(tmp_path):
    catalog = build_catalog()
    docx_path = tmp_path / "shell.docx"
    xlsx_path = tmp_path / "toc.xlsx"

    DocxShellGenerator(catalog, output_path=str(docx_path), generate_figures=False).generate()
    XlsxTocGenerator(catalog, output_path=str(xlsx_path)).generate()

    doc = Document(docx_path)
    doc_headings = [p.text for p in doc.paragraphs if p.style and p.style.name == "Heading 4"]

    workbook = load_workbook(xlsx_path, read_only=True, data_only=True)
    master_sheet = workbook["TOC_Master"]
    workbook_ids = [row[0] for row in master_sheet.iter_rows(min_row=2, values_only=True)]

    for item_id in REMOVED_IDS:
        assert item_id not in workbook_ids

    for heading in REMOVED_HEADINGS:
        assert heading not in doc_headings
