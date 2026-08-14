from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import RGBColor

from tflshell.docx_utils.header_block import add_tfl_page_header
from tflshell.models.enums import Section, TFLType
from tflshell.models.tfl_item import TFLItem


def _item() -> TFLItem:
    return TFLItem(
        id="T14.1.99",
        title="Example Clinical Table",
        tfl_type=TFLType.TABLE,
        section=Section.SEC_14_1,
        population="Safety Population",
    )


def test_header_has_one_black_bold_title_with_bookmark():
    doc = Document()
    item = _item()

    add_tfl_page_header(doc, item)

    headings = [p for p in doc.paragraphs if p.style.name == "Heading 4"]
    assert len(headings) == 1
    assert headings[0].text == f"{item.display_label}  {item.title}"
    title_runs = [run for run in headings[0].runs if run.text]
    assert len(title_runs) == 1
    assert title_runs[0].bold is True
    assert title_runs[0].italic is False
    assert title_runs[0].font.color.rgb == RGBColor(0, 0, 0)

    bookmarks = headings[0]._p.xpath(".//w:bookmarkStart")
    assert len(bookmarks) == 1
    assert bookmarks[0].get(qn("w:name")) == item.id
    assert sum(item.title in p.text for p in doc.paragraphs) == 1


def test_sponsor_and_page_share_first_line_with_right_tab():
    doc = Document()
    item = _item()

    add_tfl_page_header(doc, item)

    first_line = doc.paragraphs[0]
    assert first_line.text.startswith(f"Sponsor: {item.sponsor_placeholder}")
    assert "Page  of " in first_line.text
    instructions = [node.text for node in first_line._p.xpath(".//w:instrText")]
    assert instructions == ["PAGE", "NUMPAGES"]
    tab_stops = list(first_line.paragraph_format.tab_stops)
    assert len(tab_stops) == 1
    assert tab_stops[0].alignment == WD_TAB_ALIGNMENT.RIGHT
    assert doc.paragraphs[1].text == f"Protocol: {item.protocol_placeholder}"
