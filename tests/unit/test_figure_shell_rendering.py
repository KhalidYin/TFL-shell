import matplotlib.pyplot as plt
from docx import Document

from tflshell.data.definitions import build_catalog
from tflshell.figures.registry import (
    FIGURE_CLASS_MAP,
    build_mock_figure_data,
    generate_figure_buffer,
    supported_figure_types,
)
from tflshell.generators.docx_shell import DocxShellGenerator
from tflshell.models.catalog import TFLCatalog


def test_all_catalog_figure_types_have_registry_renderer():
    catalog = build_catalog()
    figure_types = {item.figure_type for item in catalog.figures() if item.figure_type}

    assert figure_types <= supported_figure_types()


def test_representative_figure_types_generate_non_empty_png_buffers():
    catalog = build_catalog()
    representatives = {}
    for item in catalog.figures():
        representatives.setdefault(item.figure_type, item)

    assert representatives
    for item in representatives.values():
        buffer = generate_figure_buffer(item)
        assert buffer.getbuffer().nbytes > 10_000, item.id


def test_high_priority_figures_use_dedicated_semantic_renderers():
    expected_classes = {
        "cdf": "CDFFigure",
        "edish": "EDISHFigure",
        "lab_toxicity_heatmap": "LabToxicityHeatmapFigure",
        "concentration_qtc": "ConcentrationQTcFigure",
        "pk_profile": "PKProfileFigure",
        "food_effect_profile": "PKProfileFigure",
    }

    assert {name: FIGURE_CLASS_MAP[name].__name__ for name in expected_classes} == expected_classes


def test_high_priority_figure_profiles_have_clinically_meaningful_axes_and_marks():
    expectations = {
        "cdf": ("Change from Baseline", "Cumulative Probability"),
        "edish": ("Peak ALT or AST", "Peak Total Bilirubin"),
        "lab_toxicity_heatmap": ("Laboratory Parameter", "Subjects"),
        "concentration_qtc": ("Plasma Drug Concentration", "QTcF"),
        "pk_profile": ("Nominal Time After Dose", "Mean Drug Concentration"),
    }

    for figure_type, (expected_x, expected_y) in expectations.items():
        figure = FIGURE_CLASS_MAP[figure_type]().build(build_mock_figure_data(figure_type))
        axis = figure.axes[0]
        assert expected_x in axis.get_xlabel(), figure_type
        assert expected_y in axis.get_ylabel(), figure_type
        if figure_type == "edish":
            assert axis.get_xscale() == "log"
            assert axis.get_yscale() == "log"
            assert len(axis.collections) >= 1
        elif figure_type == "lab_toxicity_heatmap":
            assert len(axis.images) == 1
        else:
            assert axis.lines or axis.collections
        plt.close(figure)


def test_mock_figure_profiles_are_explicitly_marked_as_shell_data():
    for figure_type in supported_figure_types():
        assert build_mock_figure_data(figure_type)["_shell_mock"] is True


def test_docx_shell_embeds_generated_figure_image(tmp_path):
    figure_item = build_catalog().get("F14.2.4")
    output_path = tmp_path / "figure_shell.docx"

    DocxShellGenerator(
        TFLCatalog([figure_item]),
        output_path=str(output_path),
        generate_figures=True,
    ).generate()

    doc = Document(output_path)
    paragraphs = "\n".join(paragraph.text for paragraph in doc.paragraphs)

    assert len(doc.inline_shapes) >= 1
    assert "INSERT FIGURE HERE" not in paragraphs


def test_docx_generator_removes_only_trailing_tfl_page_break(tmp_path):
    catalog = build_catalog()
    output_path = tmp_path / "two_shells.docx"
    items = [catalog.get("F14.2.3"), catalog.get("F14.3.2.1")]

    DocxShellGenerator(
        TFLCatalog(items),
        output_path=str(output_path),
        generate_figures=False,
    ).generate()

    doc = Document(output_path)
    final_paragraph = doc.paragraphs[-1]
    assert not final_paragraph._p.xpath('.//w:br[@w:type="page"]')
    assert sum(
        bool(paragraph._p.xpath('.//w:br[@w:type="page"]')) for paragraph in doc.paragraphs
    ) >= 1
