from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from tflshell.docx_utils.three_line_table import create_three_line_table
from tflshell.models.enums import Section, TFLType
from tflshell.models.tfl_item import TFLItem

MODEL_HEADER = [
    [
        {"label": "Visit / Statistic", "rowspan": 2, "alignment": "left"},
        {"label": "Treatment Estimates", "colspan": 2},
        {"label": "Treatment Comparison", "colspan": 2},
    ],
    [
        {"label": "Group 1\n(N=XX)"},
        {"label": "Group 2\n(N=XX)"},
        {"label": "Group 1 vs Group 2\nEstimate (SE)"},
        {"label": "95% CI / p-value"},
    ],
]


def _item():
    return TFLItem(
        id="T14.2.99",
        title="Model-based Change from Baseline by Visit",
        tfl_type=TFLType.TABLE,
        section=Section.SEC_14_2,
        population="Full Analysis Set",
        placeholder_columns=["Visit / Statistic", "Group 1", "Group 2", "Comparison", "CI"],
        header_rows=MODEL_HEADER,
        layout_profile="model-comparison",
        column_alignments=["left", "center", "center", "center", "center"],
        shell_rows=[
            {"label": "Week 12 — LS Mean (SE)", "values": ["xx.x", "xx.x", "xx.x", "[xx.x, xx.x]"]}
        ],
    )


def test_layout_contract_keeps_legacy_columns_and_declares_comparison_group():
    item = _item()

    assert item.leaf_column_count == 5
    assert item.effective_header_rows == MODEL_HEADER
    assert item.comparison_position == "Independent Treatment Comparison column group"


def test_multilevel_header_renders_spans_and_repeating_rows():
    item = _item()
    doc = Document()
    table = create_three_line_table(
        doc,
        rows=2,
        cols=item.leaf_column_count,
        headers=item.placeholder_columns,
        header_rows=item.effective_header_rows,
        column_alignments=item.column_alignments,
        data_rows=item.shell_data_rows_rich,
    )

    assert len(table.rows) == 3
    assert table.cell(0, 1).text == "Treatment Estimates"
    assert table.cell(0, 3).text == "Treatment Comparison"
    assert "Group 1 vs Group 2" in table.cell(1, 3).text
    assert all(
        paragraph.alignment == WD_ALIGN_PARAGRAPH.LEFT for paragraph in table.cell(0, 0).paragraphs
    )
    assert table.cell(2, 0).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert table.cell(2, 1).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    for row in table.rows[:2]:
        assert row._tr.get_or_add_trPr().find(qn("w:tblHeader")) is not None


def test_flat_header_remains_backward_compatible():
    doc = Document()
    table = create_three_line_table(
        doc,
        rows=2,
        cols=3,
        headers=["Parameter", "Group 1", "Group 2"],
        data_rows=[["Mean", "xx.x", "xx.x"]],
    )

    assert len(table.rows) == 2
    assert [cell.text for cell in table.rows[0].cells] == ["Parameter", "Group 1", "Group 2"]
    assert [cell.text for cell in table.rows[1].cells] == ["Mean", "xx.x", "xx.x"]
    assert table.cell(0, 0).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert table.cell(0, 1).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
